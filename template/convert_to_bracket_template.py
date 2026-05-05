#!/usr/bin/env python
"""
Template Bracket Converter Script

This script converts a standard Word/Excel template to use bracket-style placeholders.
It replaces standard placeholder texts with [Placeholder] format.
"""

import os
import sys
import re
from pathlib import Path
from docx import Document
from openpyxl import load_workbook

def convert_docx_to_bracket_template(input_path, output_path):
    """
    Convert a DOCX template to use bracket-style placeholders.

    Args:
        input_path: Path to the input DOCX file
        output_path: Path to save the converted file
    """
    print(f"Converting {input_path} to bracket template...")

    # Load the document
    doc = Document(input_path)

    # Placeholder replacements
    replacements = {
        "Company Name": "[Company Name]",
        "Your Company Slogan": "[Company Slogan]",
        "Street Address": "[Company Address]",
        "City, ST ZIP Code": "[Company City]",
        "Phone": "[Company Phone]",
        "555-1234": "[Company Phone]",
        "[555-1234]": "[Company Phone]",
        "Email": "[Company Email]",
        "Website": "[Company Website]",

        "INVOICE #": "INVOICE #: [Invoice Number]",
        "DATE": "DATE: [Invoice Date]",
        "DUE DATE": "DUE DATE: [Due Date]",
        "TERMS": "TERMS: [Payment Terms]",

        "Recipient Name": "[Client Name]",
        "Company Name": "[Client Company]",
        "Street Address": "[Client Address]",
        "City, ST ZIP Code": "[Client City]",
        "Phone": "[Client Phone]",

        "SUBTOTAL": "SUBTOTAL: [Subtotal]",
        "DISCOUNT": "DISCOUNT: [Discount Amount]",
        "TAX": "TAX: [Tax Amount]",
        "TOTAL": "TOTAL: [Total Amount]",
        "BALANCE DUE": "BALANCE DUE: [Total Amount]"
    }

    # Process paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text:
            original_text = paragraph.text
            new_text = original_text

            # Apply replacements
            for old_text, new_text_template in replacements.items():
                if old_text in new_text:
                    new_text = new_text.replace(old_text, new_text_template)

            # Update paragraph if changes were made
            if new_text != original_text:
                print(f"Replacing: '{original_text}' -> '{new_text}'")
                paragraph.text = new_text

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text:
                        original_text = paragraph.text
                        new_text = original_text

                        # Apply replacements
                        for old_text, new_text_template in replacements.items():
                            if old_text in new_text:
                                new_text = new_text.replace(old_text, new_text_template)

                        # Update paragraph if changes were made
                        if new_text != original_text:
                            print(f"Replacing in table: '{original_text}' -> '{new_text}'")
                            paragraph.text = new_text

    # Save the modified document
    doc.save(output_path)
    print(f"Saved bracket template to {output_path}")

def convert_xlsx_to_bracket_template(input_path, output_path):
    """
    Convert an XLSX template to use bracket-style placeholders.

    Args:
        input_path: Path to the input XLSX file
        output_path: Path to save the converted file
    """
    print(f"Converting {input_path} to bracket template...")

    # Load workbook
    wb = load_workbook(input_path)

    # Placeholder replacements
    replacements = {
        "Company Name": "[Company Name]",
        "Your Company Slogan": "[Company Slogan]",
        "Street Address": "[Company Address]",
        "City, ST ZIP Code": "[Company City]",
        "Phone": "[Company Phone]",
        "555-1234": "[Company Phone]",
        "Email": "[Company Email]",
        "Website": "[Company Website]",

        "INVOICE #": "INVOICE #: [Invoice Number]",
        "DATE": "DATE: [Invoice Date]",
        "DUE DATE": "DUE DATE: [Due Date]",
        "TERMS": "TERMS: [Payment Terms]",

        "Recipient Name": "[Client Name]",
        "Company Name": "[Client Company]",
        "Street Address": "[Client Address]",
        "City, ST ZIP Code": "[Client City]",
        "Phone": "[Client Phone]",

        "SUBTOTAL": "SUBTOTAL: [Subtotal]",
        "DISCOUNT": "DISCOUNT: [Discount Amount]",
        "TAX": "TAX: [Tax Amount]",
        "TOTAL": "TOTAL: [Total Amount]",
        "BALANCE DUE": "BALANCE DUE: [Total Amount]"
    }

    # For each worksheet, replace placeholders
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]

        # Process each cell
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    original_text = cell.value
                    new_text = original_text

                    # Apply replacements
                    for old_text, new_text_template in replacements.items():
                        if old_text in new_text:
                            new_text = new_text.replace(old_text, new_text_template)

                    # Update cell if changes were made
                    if new_text != original_text:
                        print(f"Replacing in {sheet_name}: '{original_text}' -> '{new_text}'")
                        cell.value = new_text

    # Save the modified workbook
    wb.save(output_path)
    print(f"Saved bracket template to {output_path}")

def main():
    if len(sys.argv) < 2:
        print("Usage: python convert_to_bracket_template.py <input_template> [output_template]")
        print("If output_template is not specified, 'Bracket' will be prepended to the filename.")
        sys.exit(1)

    input_path = sys.argv[1]

    # Determine output path
    if len(sys.argv) >= 3:
        output_path = sys.argv[2]
    else:
        input_file = Path(input_path)
        output_dir = input_file.parent
        output_name = f"Bracket{input_file.name}"
        output_path = output_dir / output_name

    # Ensure input file exists
    if not os.path.exists(input_path):
        print(f"Error: Input file {input_path} does not exist")
        sys.exit(1)

    # Convert based on file type
    if input_path.lower().endswith('.docx'):
        convert_docx_to_bracket_template(input_path, output_path)
    elif input_path.lower().endswith('.xlsx'):
        convert_xlsx_to_bracket_template(input_path, output_path)
    else:
        print(f"Error: Unsupported file format. Only DOCX and XLSX are supported.")
        sys.exit(1)

if __name__ == "__main__":
    main()