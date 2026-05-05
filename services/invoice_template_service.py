#!/usr/bin/env python
"""
Invoice Template Service

This module handles the selection, population, and generation of invoices
using templates based on the user's request type.
"""

import os
import logging
import tempfile
import uuid
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime, date
import shutil
import json
import re
import sys
import subprocess

# For Excel templates
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter

# For Word templates
from docx import Document
from docx.table import Table

# For PDF conversion
try:
    import pythoncom
    import win32com.client
    HAS_WIN32 = True
except ImportError:
    HAS_WIN32 = False

try:
    import uno
    import unohelper
    from com.sun.star.beans import PropertyValue
    HAS_LIBREOFFICE = True
except ImportError:
    HAS_LIBREOFFICE = False

# For XML parsing in Word templates
from xml.etree import ElementTree as ET
try:
    from lxml import etree
    HAS_LXML = True
except ImportError:
    HAS_LXML = False

# Import constants instead of hardcoding them
from constants.invoice_template_constants import (
    BASE_DIR,
    TEMPLATE_DIR,
    OUTPUT_DIR,
    DEFAULT_TEMPLATE,
    TEMPLATE_TYPES,
    FIELD_MAPPINGS,
    DIRECT_REPLACEMENTS,
    FIELD_VARIATIONS,
    IMPORTANT_TEMPLATE_FIELDS,
    CURRENCY_SYMBOLS,
    MONETARY_FIELDS
)

logger = logging.getLogger(__name__)

# Ensure the output directory exists
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Define XML parsing function for Word document styling
def parse_xml(xml_string):
    """
    Parse XML string to XML element for use in docx styling.

    Args:
        xml_string: XML string to parse

    Returns:
        XML element
    """
    if HAS_LXML:
        return etree.fromstring(xml_string)
    else:
        return ET.fromstring(xml_string)

def select_template(invoice_data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Select the appropriate template based on invoice data.

    Args:
        invoice_data: Dictionary containing invoice information

    Returns:
        Dict containing template selection results
    """
    result = {
        'success': True,
        'template_path': None,
        'template_type': None,
        'template_info': None,
        'missing_fields': []
    }

    # Default to professional template
    template_key = "default"
    best_confidence = 0

    # Get the invoice type if specified
    invoice_type = invoice_data.get("invoice_type", "").lower()

    # Try to match based on invoice type or keywords in the data
    if invoice_type:
        # Try to directly match the invoice type to a template key
        for key, info in TEMPLATE_TYPES.items():
            # Direct match with template key
            if invoice_type in key.lower():
                template_key = key
                best_confidence = 1.0
                logger.info(f"Found direct template match: {key} for invoice_type={invoice_type}")
                break

            # Match with keywords
            for keyword in info["keywords"]:
                if keyword.lower() in invoice_type:
                    # If this is a better match than what we've found so far
                    confidence = 0.8
                    if confidence > best_confidence:
                        template_key = key
                        best_confidence = confidence
                        logger.info(f"Found keyword match: {key} (keyword={keyword}) for invoice_type={invoice_type}")

    # If no match found, use default
    if best_confidence == 0:
        template_key = "default"
        best_confidence = 0.5  # Medium confidence
        logger.info(f"No template match found, using default template: {template_key}")

    # Get template info
    template_info = TEMPLATE_TYPES.get(template_key, TEMPLATE_TYPES["default"])

    # Check for missing required fields
    missing_fields = check_missing_fields(template_info, invoice_data)

    # Set the template file path
    template_file = template_info.get('file', DEFAULT_TEMPLATE)
    template_path = TEMPLATE_DIR / template_file

    # Update result
    result['template_path'] = template_path
    result['template_type'] = template_key
    result['template_info'] = template_info
    result['missing_fields'] = missing_fields

    logger.info(f"Selected template: {template_key} ({template_info['name']}) with confidence {best_confidence:.2f}")
    if missing_fields:
        logger.warning(f"Missing fields for template {template_key}: {', '.join(missing_fields)}")

    return result

def check_missing_fields(template_info: Dict[str, Any], invoice_data: Dict[str, Any]) -> List[str]:
    """
    Check which required fields are missing from the invoice data.

    Args:
        template_info: The template information
        invoice_data: The invoice data to check

    Returns:
        List of missing field names
    """
    missing_fields = []

    for field in template_info["required_fields"]:
        # Special handling for nested fields
        if field == "items":
            if "items" not in invoice_data or not invoice_data["items"]:
                missing_fields.append("items (at least one item required)")
        elif field == "labor_items":
            if "labor_items" not in invoice_data or not invoice_data["labor_items"]:
                missing_fields.append("labor_items (at least one labor entry required)")
        elif field == "materials":
            if "materials" not in invoice_data or not invoice_data["materials"]:
                missing_fields.append("materials (at least one material entry required)")
        # Regular fields
        elif field not in invoice_data or not invoice_data[field]:
            # Format field name for better readability
            formatted_field = field.replace("_", " ").title()
            missing_fields.append(formatted_field)

    return missing_fields

def populate_docx_template(template_file: str, invoice_data: Dict[str, Any]) -> str:
    """
    Populate a DOCX template with invoice data.

    Args:
        template_file: The template file path
        invoice_data: The invoice data to populate

    Returns:
        Path to the generated document
    """
    # Create output filename
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    filename = f"invoice_{invoice_data.get('invoice_number', 'INV-' + uuid.uuid4().hex[:8])}_{timestamp}.docx"
    output_path = OUTPUT_DIR / filename
    os.makedirs(output_path.parent, exist_ok=True)

    # Load template
    template_path = TEMPLATE_DIR / template_file
    logger.info(f"Loading template from: {template_path}")
    doc = Document(template_path)

    # Extract and log all placeholder texts from the document to better understand the template
    all_placeholder_texts = []
    logger.info("Analyzing document to identify placeholders:")
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text:
            # Log first 50 chars of each paragraph
            logger.info(f"Paragraph {i}: {paragraph.text[:70]}")

            # Look for potential placeholders using various patterns
            bracket_matches = re.findall(r'\[(.*?)\]', paragraph.text)
            if bracket_matches:
                logger.info(f"  Found bracketed placeholders in paragraph {i}: {bracket_matches}")
                all_placeholder_texts.extend(bracket_matches)

    # Log summary of found placeholders
    logger.info(f"All detected placeholders in document: {set(all_placeholder_texts)}")

    # Log all fields we're trying to fill
    logger.info(f"Fields available for template population: {list(invoice_data.keys())}")

    # Process paragraphs - use our new function for bracket placeholders
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text:
            # Process all bracket-style placeholders in the paragraph
            original_text = paragraph.text
            new_text = process_bracket_placeholders(original_text, invoice_data)

            if new_text != original_text:
                # Update paragraph text while preserving formatting if possible
                if paragraph.runs:
                    # If there are multiple runs with different formatting,
                    # we need to be careful to preserve the formatting
                    for run in paragraph.runs:
                        if run.text:
                            run_text = process_bracket_placeholders(run.text, invoice_data)
                            if run_text != run.text:
                                run.text = run_text
                        else:
                            # Single run paragraph, just update the text
                            paragraph.text = new_text

    # Process tables - handle both header cells and item rows
    for t, table in enumerate(doc.tables):
        logger.info(f"Processing table {t} with {len(table.rows)} rows")

        # Check if this could be the items table
        has_item_headers = False
        header_row_idx = -1

        # Look for common item table headers in all rows
        for i in range(min(3, len(table.rows))):  # Check first 3 rows
            row_text = " ".join(cell.text.upper() for cell in table.rows[i].cells)
            if any(header in row_text for header in ["QUANTITY", "DESCRIPTION", "AMOUNT", "PRICE", "TOTAL", "QTY", "RATE"]):
                has_item_headers = True
                header_row_idx = i
                logger.info(f"Found items table at table {t}, header row {i}: '{row_text}'")
                break

        # Process each cell in the table
        for r, row in enumerate(table.rows):
            for c, cell in enumerate(row.cells):
                # Process all text in the cell for bracket placeholders
                if cell.text:
                    original_text = cell.text
                    new_text = process_bracket_placeholders(original_text, invoice_data)

                    if new_text != original_text:
                        # Update cell text
                        for paragraph in cell.paragraphs:
                            paragraph.text = process_bracket_placeholders(paragraph.text, invoice_data)

        # If this looks like an items table, populate it with invoice items
        if has_item_headers and "items" in invoice_data and invoice_data["items"]:
            # Find column indexes for item fields
            header_row = table.rows[header_row_idx]
            header_texts = [cell.text.upper() for cell in header_row.cells]

            # Determine column indexes
            desc_col = next((i for i, text in enumerate(header_texts)
                           if "DESCRIPTION" in text or "ITEM" in text or "SERVICE" in text), -1)
            qty_col = next((i for i, text in enumerate(header_texts)
                          if "QUANTITY" in text or "QTY" in text), -1)
            rate_col = next((i for i, text in enumerate(header_texts)
                           if "RATE" in text or "PRICE" in text or "UNIT" in text), -1)
            amount_col = next((i for i, text in enumerate(header_texts)
                             if "AMOUNT" in text or "TOTAL" in text), -1)

            logger.info(f"Item table columns - Description: {desc_col}, Quantity: {qty_col}, Rate: {rate_col}, Amount: {amount_col}")

            # Skip if we couldn't identify the columns
            if desc_col == -1:
                logger.warning("Could not identify item description column, skipping item population")
                continue

            # Find where to start inserting items (below header row)
            start_row = header_row_idx + 1
            items = invoice_data["items"]

            # Clear any existing content in item rows
            for r in range(start_row, min(start_row + len(items), len(table.rows))):
                for c in range(len(table.rows[r].cells)):
                    # Clear cell content but preserve formatting
                    if table.rows[r].cells[c].paragraphs:
                        table.rows[r].cells[c].paragraphs[0].text = ""

            # Populate items
            for i, item in enumerate(items):
                # Make sure we have enough rows
                if start_row + i >= len(table.rows):
                    # Need to add a new row
                    table.add_row()

                row = table.rows[start_row + i]

                # Populate the cells
                if desc_col >= 0 and len(row.cells) > desc_col:
                    row.cells[desc_col].text = str(item.get("description", ""))

                if qty_col >= 0 and len(row.cells) > qty_col:
                    qty = item.get("quantity", 0)
                    row.cells[qty_col].text = str(qty)

                if rate_col >= 0 and len(row.cells) > rate_col:
                    rate = item.get("unit_price", 0)
                    currency = get_currency_symbol(invoice_data.get("currency", "USD"))
                    row.cells[rate_col].text = f"{currency}{rate:.2f}"

                if amount_col >= 0 and len(row.cells) > amount_col:
                    amount = item.get("total_price", 0)
                    currency = get_currency_symbol(invoice_data.get("currency", "USD"))
                    row.cells[amount_col].text = f"{currency}{amount:.2f}"

    # Save the document to output path
    doc.save(output_path)
    logger.info(f"Saved populated DOCX to: {output_path}")

    # Copy to UI uploads directory for easy access
    try:
        ui_uploads_dir = BASE_DIR / "ui" / "uploads"
        os.makedirs(ui_uploads_dir, exist_ok=True)

        # Create unique filename for UI
        ui_filename = f"invoice_{int(datetime.now().timestamp())}_{filename}"
        ui_path = ui_uploads_dir / ui_filename

        import shutil
        shutil.copy2(output_path, ui_path)
        logger.info(f"Copied invoice to UI uploads: {ui_path}")

        # Also set the UI path in the invoice data for the caller to use
        invoice_data["ui_doc_path"] = str(ui_path)
    except Exception as e:
        logger.warning(f"Could not copy to UI uploads: {str(e)}")

    return str(output_path)

def format_value(value, field_name):
    """
    Format a value based on its field name and type.

    Args:
        value: The value to format
        field_name: The name of the field

    Returns:
        Formatted value as a string
    """
    if value is None:
        return ""

    # Format date fields
    if isinstance(value, (datetime, date)):
        return value.strftime("%m/%d/%Y")

    # Format monetary values
    monetary_fields = ["price", "amount", "total", "cost", "subtotal", "tax", "shipping"]
    if isinstance(value, (int, float)) and any(monetary_key in field_name.lower() for monetary_key in monetary_fields):
        if field_name.lower().endswith("tax_rate") or field_name.lower().endswith("discount_rate"):
            return f"{value:.1f}%"
        else:
            return f"${value:.2f}"

    # Format numbers with commas for readability
    if isinstance(value, (int, float)) and not any(monetary_key in field_name.lower() for monetary_key in monetary_fields):
        if value == int(value):  # It's a whole number
            return f"{int(value):,}"
        else:
            return f"{value:,.2f}"

    # Default: return as string
    return str(value)

def populate_xlsx_template(template_file: str, invoice_data: Dict[str, Any]) -> str:
    """
    Populate an Excel template with invoice data.

    Args:
        template_file: The template file path
        invoice_data: The invoice data to populate

    Returns:
        Path to the generated document
    """
    # Create output filename
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    filename = f"invoice_{invoice_data.get('invoice_number', 'INV-' + uuid.uuid4().hex[:8])}_{timestamp}.xlsx"
    output_path = OUTPUT_DIR / filename
    os.makedirs(output_path.parent, exist_ok=True)

    # Get full template path
    template_path = TEMPLATE_DIR / template_file
    logger.info(f"Loading Excel template from: {template_path}")

    # Load workbook
    wb = load_workbook(template_path)
    logger.info(f"Workbook loaded with sheets: {wb.sheetnames}")

    # For each worksheet, check for placeholders
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        logger.info(f"Processing sheet: {sheet_name}")

        # Find items table (look for header rows with Description/Quantity/Rate columns)
        items_table_row = None
        items_table_cols = {}
        items_found = False

        # First identify all bracket placeholders
        bracket_placeholders = []
        for row_idx, row in enumerate(ws.rows, 1):
            for col_idx, cell in enumerate(row, 1):
                if cell.value and isinstance(cell.value, str):
                    # Look for bracket patterns [PlaceholderName]
                    matches = re.findall(r'\[(.*?)\]', cell.value)
                    if matches:
                        for match in matches:
                            bracket_placeholders.append({
                                'placeholder': f"[{match}]",
                                'cell': cell,
                                'row': row_idx,
                                'col': col_idx
                            })

        logger.info(f"Found {len(bracket_placeholders)} bracket placeholders in sheet {sheet_name}")

        # Process all bracket placeholders
        for placeholder_info in bracket_placeholders:
            cell = placeholder_info['cell']
            original_text = cell.value

            # Apply our bracket placeholder processor
            new_text = process_bracket_placeholders(original_text, invoice_data)

            if new_text != original_text:
                cell.value = new_text
                logger.info(f"Replaced placeholder in cell {get_column_letter(placeholder_info['col'])}{placeholder_info['row']}: {original_text} -> {new_text}")

        # Look for item table headers (search first 15 rows)
        for row_idx in range(1, min(15, ws.max_row + 1)):
            row_values = [ws.cell(row=row_idx, column=col_idx).value for col_idx in range(1, ws.max_column + 1)]
            row_values_str = [str(val).upper() if val else "" for val in row_values]
            row_text = " ".join(row_values_str)

            # Check if this row has item table headers
            if any(header in row_text for header in ["DESCRIPTION", "QUANTITY", "QTY", "RATE", "AMOUNT", "PRICE"]):
                items_table_row = row_idx
                logger.info(f"Found items table header at row {row_idx}: {row_text}")

                # Map columns
                for col_idx, val in enumerate(row_values_str, 1):
                    if "DESCRIPTION" in val or "ITEM" in val:
                        items_table_cols["description"] = col_idx
                    elif "QUANTITY" in val or "QTY" in val:
                        items_table_cols["quantity"] = col_idx
                    elif "RATE" in val or "PRICE" in val or "UNIT" in val:
                        items_table_cols["unit_price"] = col_idx
                    elif "AMOUNT" in val or "TOTAL" in val:
                        items_table_cols["total_price"] = col_idx

                logger.info(f"Item table columns: {items_table_cols}")

                # Only proceed if we found at least description column
                if "description" in items_table_cols:
                    items_found = True
                    break

        # Populate items if we found a table
        if items_found and "items" in invoice_data and invoice_data["items"]:
            items = invoice_data["items"]
            logger.info(f"Populating {len(items)} items in Excel template")

            # Start populating from the next row after the header
            start_row = items_table_row + 1

            # Clear any existing item rows first
            for i in range(len(items)):
                row_idx = start_row + i
                # Clear existing cells in this row
                for col_key, col_idx in items_table_cols.items():
                    ws.cell(row=row_idx, column=col_idx).value = None

            # Insert items
            for i, item in enumerate(items):
                row_idx = start_row + i

                # Description
                if "description" in items_table_cols:
                    ws.cell(row=row_idx, column=items_table_cols["description"]).value = item.get("description", "")

                # Quantity
                if "quantity" in items_table_cols:
                    ws.cell(row=row_idx, column=items_table_cols["quantity"]).value = item.get("quantity", 0)

                # Unit price
                if "unit_price" in items_table_cols:
                    ws.cell(row=row_idx, column=items_table_cols["unit_price"]).value = item.get("unit_price", 0)

                # Total price
                if "total_price" in items_table_cols:
                    ws.cell(row=row_idx, column=items_table_cols["total_price"]).value = item.get("total_price", 0)

    # Save the populated workbook
    wb.save(output_path)
    logger.info(f"Saved populated Excel file to: {output_path}")

    # Copy to UI uploads directory for easy access
    try:
        ui_uploads_dir = BASE_DIR / "ui" / "uploads"
        os.makedirs(ui_uploads_dir, exist_ok=True)

        # Create unique filename for UI
        ui_filename = f"invoice_{int(datetime.now().timestamp())}_{filename}"
        ui_path = ui_uploads_dir / ui_filename

        import shutil
        shutil.copy2(output_path, ui_path)
        logger.info(f"Copied invoice to UI uploads: {ui_path}")

        # Also set the UI path in the invoice data for the caller to use
        invoice_data["ui_doc_path"] = str(ui_path)
    except Exception as e:
        logger.warning(f"Could not copy to UI uploads: {str(e)}")

    return str(output_path)

def convert_to_pdf(input_file):
    """
    Convert a document to PDF format using available methods.

    Args:
        input_file: Path to the input document (DOCX, XLSX, etc.)

    Returns:
        str or None: Path to the generated PDF if successful, None otherwise
    """
    if not input_file or not os.path.exists(input_file):
        logger.error(f"Input file does not exist: {input_file}")
        return None

    input_ext = os.path.splitext(input_file)[1].lower()
    output_pdf = os.path.splitext(input_file)[0] + '.pdf'

    logger.info(f"Attempting to convert {input_file} to PDF: {output_pdf}")

    # If file is already PDF, just return it
    if input_ext == '.pdf':
        return input_file

    # Try different conversion methods based on file type
    if input_ext in ['.xlsx', '.xls']:
        # Excel to PDF conversion
        try:
            logger.info("Attempting conversion with win32com...")
            import win32com.client
            excel = win32com.client.Dispatch("Excel.Application")
            excel.Visible = False
            wb = excel.Workbooks.Open(os.path.abspath(input_file))
            wb.ExportAsFixedFormat(0, os.path.abspath(output_pdf))
            wb.Close()
            excel.Quit()

            if os.path.exists(output_pdf):
                logger.info(f"PDF conversion successful using win32com")
                return output_pdf
            else:
                logger.warning("PDF file not created by win32com despite no errors")
        except Exception as e:
            logger.warning(f"Failed to convert Excel to PDF using win32com: {str(e)}")

        # Try using LibreOffice if available
        try:
            logger.info("Attempting conversion with LibreOffice...")
            subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf',
                '--outdir', os.path.dirname(input_file),
                input_file
            ], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)

            if os.path.exists(output_pdf):
                logger.info(f"PDF conversion successful using LibreOffice")
                return output_pdf
            else:
                logger.warning("PDF file not created by LibreOffice despite no errors")
        except Exception as e:
            logger.warning(f"Failed to convert Excel to PDF using LibreOffice: {str(e)}")

    elif input_ext in ['.docx', '.doc']:
        # Word to PDF conversion
        try:
            logger.info("Attempting conversion with docx2pdf...")
            from docx2pdf import convert
            convert(input_file, output_pdf)

            if os.path.exists(output_pdf):
                logger.info(f"PDF conversion successful using docx2pdf")
                return output_pdf
            else:
                logger.warning("PDF file not created by docx2pdf despite no errors")
        except Exception as e:
            logger.warning(f"Failed to convert DOCX to PDF using docx2pdf: {str(e)}")

        try:
            logger.info("Attempting conversion with python-docx2pdf...")
            import aspose.words as aw
            doc = aw.Document(input_file)
            doc.save(output_pdf)

            if os.path.exists(output_pdf):
                logger.info(f"PDF conversion successful using python-docx2pdf")
                return output_pdf
            else:
                logger.warning("PDF file not created by python-docx2pdf despite no errors")
        except Exception as e:
            logger.warning(f"Failed to convert DOCX to PDF using python-docx2pdf: {str(e)}")

        try:
            logger.info("Attempting conversion with win32com...")
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(os.path.abspath(input_file))
            doc.SaveAs(os.path.abspath(output_pdf), FileFormat=17)
            doc.Close()
            word.Quit()

            if os.path.exists(output_pdf):
                logger.info(f"PDF conversion successful using win32com")
                return output_pdf
            else:
                logger.warning("PDF file not created by win32com despite no errors")
        except Exception as e:
            logger.warning(f"Failed to convert DOCX to PDF using win32com: {str(e)}")

        # Try using LibreOffice if available
        try:
            logger.info("Attempting conversion with LibreOffice...")
            subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf',
                '--outdir', os.path.dirname(input_file),
                input_file
            ], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)

            if os.path.exists(output_pdf):
                logger.info(f"PDF conversion successful using LibreOffice")
                return output_pdf
            else:
                logger.warning("PDF file not created by LibreOffice despite no errors")
        except Exception as e:
            logger.warning(f"Failed to convert DOCX to PDF using LibreOffice: {str(e)}")

        # Try pandoc as a fallback
        try:
            logger.info("Attempting conversion with pandoc...")
            subprocess.run([
                'pandoc', input_file,
                '-o', output_pdf
            ], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)

            if os.path.exists(output_pdf):
                logger.info(f"PDF conversion successful using pandoc")
                return output_pdf
            else:
                logger.warning("PDF file not created by pandoc despite no errors")
        except Exception as e:
            logger.warning(f"Failed to convert DOCX to PDF using pandoc: {str(e)}")

    # If we got here, all conversion attempts failed
    logger.error(f"All PDF conversion methods failed for {input_file}")
    return None

def generate_invoice(invoice_data, template_key=None):
    """
    Generate an invoice based on specified data.
    Always uses the default template (InvoiceDocument.docx).

    Args:
        invoice_data: Dictionary containing invoice information
        template_key: Optional parameter kept for compatibility but not used

    Returns:
        Dictionary containing result status, messages, and file paths
    """
    result = {
        'success': False,
        'message': '',
        'document_path': None,
        'pdf_path': None
    }

    try:
        # Always use the default template
        template_key = "default"
        template_info = TEMPLATE_TYPES[template_key]

        # Log the incoming invoice data for debugging
        logger.info(f"Original invoice data: {invoice_data}")

        # Check for missing required fields but continue anyway
        missing_fields = check_missing_fields(template_info, invoice_data)
        if missing_fields:
            logger.warning(f"Missing fields for template: {', '.join(missing_fields)}")

        # Ensure company profile data is properly mapped
        # Map common field variations to standard fields
        field_mappings = {
            # Company profile fields
            "company_name": ["company_name", "vendor", "business_name"],
            "company_address": ["company_address", "address", "business_address"],
            "company_phone": ["company_phone", "phone", "business_phone", "telephone"],
            "company_email": ["company_email", "email", "business_email"],
            "company_website": ["company_website", "website", "business_website"],
            "company_slogan": ["company_slogan", "tagline", "business_tagline"],
            "company_fax": ["company_fax", "fax", "business_fax"],
            "payment_terms": ["payment_terms", "terms"],

            # Client profile fields
            "client_name": ["client_name", "customer_name", "recipient_name"],
            "client_company": ["client_company", "customer_company", "recipient_company"],
            "client_address": ["client_address", "customer_address", "recipient_address"],
            "client_phone": ["client_phone", "customer_phone", "recipient_phone"],
            "client_email": ["client_email", "customer_email", "recipient_email"],
        }

        # Create a consistent data set with all possible field names
        normalized_data = {}

        # First, copy all original data
        normalized_data.update(invoice_data)

        # Then normalize all field variations
        for standard_field, variations in field_mappings.items():
            # Find the first available value among variations
            field_value = None
            for variation in variations:
                if variation in invoice_data and invoice_data[variation]:
                    field_value = invoice_data[variation]
                    break

            # If we found a value, ensure it's set for all variations
            if field_value:
                for variation in variations:
                    normalized_data[variation] = field_value

                # Log that we've set a value
                logger.info(f"Set {standard_field} to: {field_value}")

        # Replace the invoice data with our normalized version
        invoice_data = normalized_data

        # Ensure all keys from user preferences are included
        if 'preferences' in invoice_data:
            prefs = invoice_data['preferences']
            if isinstance(prefs, dict):
                for key, value in prefs.items():
                    if value:  # Only include non-empty values
                        invoice_data[key] = value
                        logger.info(f"Added preference {key}: {value}")

        # Calculate invoice totals if not provided
        if 'items' in invoice_data:
            items = invoice_data.get('items', [])

            # Calculate subtotal if needed
            if 'subtotal' not in invoice_data:
                subtotal = 0
                for item in items:
                    if 'total_price' in item:
                        subtotal += item['total_price']
                    elif 'quantity' in item and 'unit_price' in item:
                        subtotal += item['quantity'] * item['unit_price']
                invoice_data['subtotal'] = subtotal

            # Calculate tax amount if tax_rate is provided but tax_amount is not
            if 'tax_rate' in invoice_data and 'tax_amount' not in invoice_data:
                tax_rate = invoice_data.get('tax_rate', 0)
                if tax_rate:
                    subtotal = invoice_data.get('subtotal', 0)
                    tax_amount = subtotal * (tax_rate / 100)
                    invoice_data['tax_amount'] = tax_amount

            # Calculate discount amount if discount_rate is provided but discount_amount is not
            if 'discount_rate' in invoice_data and 'discount_amount' not in invoice_data:
                discount_rate = invoice_data.get('discount_rate', 0)
                if discount_rate:
                    subtotal = invoice_data.get('subtotal', 0)
                    discount_amount = subtotal * (discount_rate / 100)
                    invoice_data['discount_amount'] = discount_amount

            # Calculate total amount if needed
            if 'total_amount' not in invoice_data:
                subtotal = invoice_data.get('subtotal', 0)
                tax_amount = invoice_data.get('tax_amount', 0)
                discount_amount = invoice_data.get('discount_amount', 0)
                invoice_data['total_amount'] = subtotal + tax_amount - discount_amount

        # Get template file
        template_file = template_info.get('file')

        if not template_file:
            logger.error("No template file specified in template info")
            result['message'] = "No template file specified in template info"
            return result

        template_path = TEMPLATE_DIR / template_file

        if not os.path.exists(template_path):
            logger.error(f"Template file not found: {template_path}")
            result['message'] = f"Template file not found: {template_path}"
            return result

        logger.info(f"Using default template: {template_file}")
        logger.info(f"Final invoice data for template: {invoice_data}")

        # Populate template with data
        try:
            # Convert PosixPath to string before using string methods
            template_path_str = str(template_path)

            if template_path_str.endswith('.docx'):
                document_path = populate_docx_template(template_path, invoice_data)
            elif template_path_str.endswith(('.xlsx', '.xls')):
                document_path = populate_xlsx_template(template_path, invoice_data)
            else:
                return {
                    'success': False,
                    'message': f"Unsupported template file type: {template_path_str}",
                    'document_path': None,
                    'pdf_path': None
                }

            if not document_path:
                return {
                    'success': False,
                    'message': "Failed to generate document",
                    'document_path': None,
                    'pdf_path': None
                }

            # Try to convert to PDF
            pdf_path = convert_to_pdf(document_path)

            # Check if PDF was successfully created
            if pdf_path and os.path.exists(pdf_path):
                return {
                    'success': True,
                    'message': "Invoice generated successfully as DOCX and PDF",
                    'document_path': document_path,
                    'pdf_path': pdf_path
                }
            else:
                # PDF conversion failed but DOCX was created
                logger.warning(f"PDF conversion failed for {document_path}")
                return {
                    'success': True,
                    'message': "Invoice generated as DOCX only (PDF conversion failed)",
                    'document_path': document_path,
                    'pdf_path': None
                }
        except Exception as e:
            logger.error(f"Error generating invoice: {str(e)}")
            return {
                'success': False,
                'message': f"Error generating invoice: {str(e)}",
                'document_path': None,
                'pdf_path': None
            }

    except Exception as e:
        logger.exception(f"Error generating invoice: {str(e)}")
        result['message'] = f"Error: {str(e)}"
        return result

def generate_invoice_docx(invoice_data):
    """
    Generate invoice document based on the provided data.

    Args:
        invoice_data: Dictionary containing invoice information

    Returns:
        dict: {
            'success': bool,
            'message': str,
            'document_path': str or None,
            'pdf_path': str or None
        }
    """
    # Select appropriate template
    template_result = select_template(invoice_data)

    if not template_result['success']:
        logger.warning(f"Template selection failed: {template_result['message']}")
        return {
            'success': False,
            'message': f"Failed to select template: {template_result['message']}",
            'document_path': None,
            'pdf_path': None
        }

    template_path = template_result['template_path']
    template_type = template_result['template_type']
    missing_fields = template_result['missing_fields']

    if missing_fields:
        logger.warning(f"Missing fields for invoice: {', '.join(missing_fields)}")

    # Populate template with data
    try:
        # Convert PosixPath to string before using string methods
        template_path_str = str(template_path)

        if template_path_str.endswith('.docx'):
            document_path = populate_docx_template(template_path, invoice_data)
        elif template_path_str.endswith(('.xlsx', '.xls')):
            document_path = populate_xlsx_template(template_path, invoice_data)
        else:
            return {
                'success': False,
                'message': f"Unsupported template file type: {template_path_str}",
                'document_path': None,
                'pdf_path': None
            }

        if not document_path:
            return {
                'success': False,
                'message': "Failed to generate document",
                'document_path': None,
                'pdf_path': None
            }

        # Try to convert to PDF
        pdf_path = convert_to_pdf(document_path)

        # Check if PDF was successfully created
        if pdf_path and os.path.exists(pdf_path):
            return {
                'success': True,
                'message': "Invoice generated successfully as DOCX and PDF",
                'document_path': document_path,
                'pdf_path': pdf_path
            }
        else:
            # PDF conversion failed but DOCX was created
            logger.warning(f"PDF conversion failed for {document_path}")
            return {
                'success': True,
                'message': "Invoice generated as DOCX only (PDF conversion failed)",
                'document_path': document_path,
                'pdf_path': None
            }
    except Exception as e:
        logger.error(f"Error generating invoice: {str(e)}")
        return {
            'success': False,
            'message': f"Error generating invoice: {str(e)}",
            'document_path': None,
            'pdf_path': None
        }

def generate_invoice_xlsx(data, template_file, template_info):
    """
    Generate an invoice using an XLSX template.

    Args:
        data: Dictionary containing invoice information
        template_file: Path to the template file
        template_info: Dictionary containing template metadata

    Returns:
        Dictionary containing result and paths to generated files
    """
    result = {
        'success': False,
        'message': '',
        'document_path': None,
        'pdf_path': None
    }

    try:
        logger.info(f"Generating XLSX invoice with template: {template_file}")

        # Populate the template
        document_path = populate_xlsx_template(template_file.name, data)

        if not document_path:
            result['message'] = "Failed to populate XLSX template"
            return result

        result['document_path'] = document_path
        result['success'] = True
        result['message'] = "Invoice generated successfully"

        # Try to convert to PDF if possible
        try:
            pdf_path = convert_to_pdf(document_path)
            if pdf_path:
                result['pdf_path'] = pdf_path
                logger.info(f"Generated PDF at: {pdf_path}")
        except Exception as e:
            logger.warning(f"PDF conversion failed: {str(e)}")
            # Continue anyway since we have the XLSX

        return result

    except Exception as e:
        logger.exception(f"Error generating XLSX invoice: {str(e)}")
        result['message'] = f"Error: {str(e)}"
        return result

def generate_invoice_html(data, template_file, template_info):
    """
    Generate an invoice using an HTML template.

    Args:
        data: Dictionary containing invoice information
        template_file: Path to the template file
        template_info: Dictionary containing template metadata

    Returns:
        Dictionary containing result and paths to generated files
    """
    from datetime import timedelta
    import html
    from bs4 import BeautifulSoup

    result = {
        'success': False,
        'message': '',
        'document_path': None,
        'pdf_path': None
    }

    try:
        logger.info(f"Generating HTML invoice with template: {template_file}")

        # Read the template
        with open(template_file, 'r', encoding='utf-8') as f:
            template_content = f.read()

        # Prepare template data
        template_data = data.copy()

        # Ensure we have default values for common fields
        if 'invoice_date' not in template_data or not template_data['invoice_date']:
            template_data['invoice_date'] = datetime.now().strftime('%Y-%m-%d')

        if 'due_date' not in template_data or not template_data['due_date']:
            # Default to 30 days from invoice date
            due_date = datetime.now() + timedelta(days=30)
            template_data['due_date'] = due_date.strftime('%Y-%m-%d')

        if 'status' not in template_data or not template_data['status']:
            template_data['status'] = 'Pending'

        # Format currency properly
        if 'currency' not in template_data:
            template_data['currency'] = 'USD'

        template_data['currency_symbol'] = get_currency_symbol(template_data['currency'])

        # Make sure client information is available
        client_fields = ['client_name', 'client_company', 'client_address', 'client_phone', 'client_email']
        for field in client_fields:
            if field not in template_data or not template_data[field]:
                if field == 'client_name' and 'vendor' in template_data and template_data['vendor'] != template_data.get('company_name'):
                    template_data[field] = template_data['vendor']
                    logger.info(f"Using vendor as client_name: {template_data['vendor']}")
                else:
                    template_data[field] = ""

        # Make sure items is always a list
        if 'items' not in template_data:
            template_data['items'] = []

        items = template_data['items']

        # Calculate totals if needed
        if items:
            # Calculate subtotal if needed
            if 'subtotal' not in template_data:
                subtotal = sum(item.get('total_price', 0) for item in items)
                template_data['subtotal'] = subtotal

            # Ensure tax and discount exist
            template_data.setdefault('tax_rate', 0)
            template_data.setdefault('discount_rate', 0)
            template_data.setdefault('tax_amount', 0)
            template_data.setdefault('discount_amount', 0)

        # Prepare HTML for items table
        soup = BeautifulSoup(template_content, 'html.parser')
        items_table = soup.find('table')

        if items_table and items:
            # Find and clear the sample row
            tbody = items_table.find('tbody')
            if tbody:
                tbody.clear()
            else:
                tbody = soup.new_tag('tbody')
                items_table.append(tbody)

            # Add rows for each item
            for i, item in enumerate(items, 1):
                tr = soup.new_tag('tr')
                # Add alternating row colors
                if i % 2 == 0:
                    tr['style'] = 'background-color: #f9f9f9;'

                # Item number
                td_item = soup.new_tag('td')
                td_item.string = str(i)
                tr.append(td_item)

                # Description
                td_desc = soup.new_tag('td')
                td_desc.string = item.get('description', '')
                tr.append(td_desc)

                # Quantity
                td_qty = soup.new_tag('td', **{'class': 'text-center'})
                td_qty.string = str(item.get('quantity', ''))
                tr.append(td_qty)

                # Unit price
                td_price = soup.new_tag('td', **{'class': 'text-right'})
                unit_price = item.get('unit_price', 0)
                td_price.string = f"{template_data['currency_symbol']}{unit_price:.2f}"
                tr.append(td_price)

                # Total amount
                td_amount = soup.new_tag('td', **{'class': 'text-right'})
                total_price = item.get('total_price', unit_price * item.get('quantity', 0))
                td_amount.string = f"{template_data['currency_symbol']}{total_price:.2f}"
                tr.append(td_amount)

                tbody.append(tr)

            # Update the template content with the modified HTML
            template_content = str(soup)

        # Apply template variables
        for key, value in template_data.items():
            if key != 'items':  # Skip items array which is handled separately
                # Format numeric values
                if isinstance(value, (int, float)):
                    if key in ['tax_rate', 'discount_rate']:
                        value = f"{value:.1f}"
                    elif key in ['subtotal', 'tax_amount', 'discount_amount', 'total_amount']:
                        value = f"{value:.2f}"
                    else:
                        value = str(value)
                else:
                    value = str(value)

                value = html.escape(value)

                # Replace {{key}} format
                placeholder = '{{' + key + '}}'
                template_content = template_content.replace(placeholder, value)

                # Replace [key] format
                bracket_placeholder = '[' + key + ']'
                template_content = template_content.replace(bracket_placeholder, value)

        # Generate a unique filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        invoice_number = template_data.get('invoice_number', '').replace(' ', '_')
        if not invoice_number:
            invoice_number = f"invoice_{uuid.uuid4().hex[:8]}"

        output_basename = f"invoice_{timestamp}_{invoice_number}"
        html_path = OUTPUT_DIR / f"{output_basename}.html"

        # Ensure output directory exists
        os.makedirs(OUTPUT_DIR, exist_ok=True)

        # Save HTML file
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(template_content)

        logger.info(f"HTML invoice saved to {html_path}")

        # NOTE: PDF conversion would normally happen here
        # For this implementation, we're just using the HTML file

        result['success'] = True
        result['document_path'] = str(html_path)
        result['message'] = "Invoice generated successfully"

        return result

    except Exception as e:
        logger.exception(f"Error generating HTML invoice: {str(e)}")
        result['message'] = f"Error: {str(e)}"
        return result

def generate_basic_html_template(template_data):
    """Generate a basic HTML template using the provided data"""
    company_name = template_data.get('company_name', 'Your Company')
    company_address = template_data.get('company_address', '123 Company St')
    company_phone = template_data.get('company_phone', '555-1234')
    company_email = template_data.get('company_email', 'company@example.com')

    client_name = template_data.get('client_name', '')
    client_address = template_data.get('client_address', '')
    client_phone = template_data.get('client_phone', '')
    client_email = template_data.get('client_email', '')

    invoice_number = template_data.get('invoice_number', 'INV-123')
    invoice_date = template_data.get('invoice_date', datetime.now().strftime('%Y-%m-%d'))
    due_date = template_data.get('due_date', '')

    items = template_data.get('items', [])
    total_amount = template_data.get('total_amount', 0)
    currency = template_data.get('currency', 'USD')
    currency_symbol = get_currency_symbol(currency)

    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Invoice {invoice_number}</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 0; padding: 20px; color: #333; }}
            .invoice-header {{ margin-bottom: 20px; }}
            .company-details {{ margin-bottom: 20px; }}
            .client-details {{ margin-bottom: 20px; }}
            .invoice-details {{ margin-bottom: 20px; }}
            table {{ width: 100%; border-collapse: collapse; margin-bottom: 20px; }}
            th, td {{ padding: 10px; text-align: left; border-bottom: 1px solid #ddd; }}
            th {{ background-color: #f2f2f2; }}
            .total {{ font-weight: bold; text-align: right; margin-top: 20px; }}
            h1 {{ color: #2a5885; }}
        </style>
    </head>
    <body>
        <div class="invoice-header">
            <h1>INVOICE</h1>
        </div>

        <div class="company-details">
            <h2>{company_name}</h2>
            <p>{company_address}</p>
            <p>Phone: {company_phone}</p>
            <p>Email: {company_email}</p>
        </div>

        <div class="client-details">
            <h3>Bill To:</h3>
            <p>{client_name}</p>
            <p>{client_address}</p>
            <p>Phone: {client_phone}</p>
            <p>Email: {client_email}</p>
        </div>

        <div class="invoice-details">
            <p><strong>Invoice Number:</strong> {invoice_number}</p>
            <p><strong>Invoice Date:</strong> {invoice_date}</p>
            <p><strong>Due Date:</strong> {due_date}</p>
        </div>

        <table>
            <thead>
                <tr>
                    <th>Description</th>
                    <th>Quantity</th>
                    <th>Unit Price</th>
                    <th>Total</th>
                </tr>
            </thead>
            <tbody>
    """

    for item in items:
        description = item.get('description', '')
        quantity = item.get('quantity', 1)
        unit_price = item.get('unit_price', 0)
        total_price = item.get('total_price', 0)

        html += f"""
                <tr>
                    <td>{description}</td>
                    <td>{quantity}</td>
                    <td>{currency_symbol}{unit_price}</td>
                    <td>{currency_symbol}{total_price}</td>
                </tr>
        """

    html += f"""
            </tbody>
        </table>

        <div class="total">
            <p>Total: {currency_symbol}{total_amount} {currency}</p>
        </div>

        <div class="notes">
            <p><strong>Notes:</strong></p>
            <p>Thank you for your business!</p>
        </div>
    </body>
    </html>
    """

    return html

def generate_basic_pdf(template_data, output_path):
    """Generate a basic PDF using reportlab if available"""
    try:
        # First, check if reportlab is installed
        import importlib
        reportlab_spec = importlib.util.find_spec("reportlab")
        if reportlab_spec is None:
            logger.warning("ReportLab library not installed. Cannot generate PDF.")
            return None

        # If reportlab is available, import the necessary modules
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet

        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []

        # Company info
        company_name = template_data.get('company_name', 'Your Company')
        company_address = template_data.get('company_address', '')

        # Invoice details
        invoice_number = template_data.get('invoice_number', 'INV-123')
        invoice_date = template_data.get('invoice_date', datetime.now().strftime('%Y-%m-%d'))

        # Client info
        client_name = template_data.get('client_name', '')

        # Add invoice title
        elements.append(Paragraph(f"INVOICE #{invoice_number}", styles['Title']))
        elements.append(Spacer(1, 12))

        # Add company info
        elements.append(Paragraph(f"From: {company_name}", styles['Heading2']))
        elements.append(Paragraph(company_address, styles['Normal']))
        elements.append(Spacer(1, 12))

        # Add client info
        elements.append(Paragraph(f"To: {client_name}", styles['Heading2']))
        elements.append(Spacer(1, 12))

        # Add invoice date
        elements.append(Paragraph(f"Date: {invoice_date}", styles['Normal']))
        elements.append(Spacer(1, 24))

        # Add items table
        items = template_data.get('items', [])
        items_data = [['Description', 'Quantity', 'Unit Price', 'Total']]

        for item in items:
            description = item.get('description', '')
            quantity = item.get('quantity', 1)
            unit_price = item.get('unit_price', 0)
            total_price = item.get('total_price', 0)
            items_data.append([description, str(quantity), str(unit_price), str(total_price)])

        # Add total row
        total_amount = template_data.get('total_amount', 0)
        currency = template_data.get('currency', 'USD')
        items_data.append(['', '', 'Total:', f"{total_amount} {currency}"])

        # Create table
        table = Table(items_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))

        elements.append(table)
        elements.append(Spacer(1, 36))
        elements.append(Paragraph("Thank you for your business!", styles['Normal']))

        # Build PDF
        doc.build(elements)
        return output_path
    except ImportError:
        logger.warning("ReportLab not available, cannot generate basic PDF")
        return None
    except Exception as e:
        logger.error(f"Error generating PDF with ReportLab: {str(e)}")
        return None

def get_currency_symbol(currency_code):
    """Get the currency symbol for a currency code"""
    symbols = {
        'USD': '$',
        'EUR': '€',
        'GBP': '£',
        'JPY': '¥',
        'INR': '₹',
        'CAD': 'C$',
        'AUD': 'A$',
        'CHF': 'CHF',
        'CNY': '¥',
        'NZD': 'NZ$',
    }
    return symbols.get(currency_code, currency_code)

def process_bracket_placeholders(text: str, invoice_data: Dict[str, Any]) -> str:
    """
    Process all bracket-style placeholders in text and replace with invoice data.

    Args:
        text: Text containing [Placeholder] style placeholders
        invoice_data: Dictionary containing invoice data fields

    Returns:
        Text with placeholders replaced by actual values
    """
    if not text:
        return text

    # Find all [Placeholder] patterns
    bracket_matches = re.findall(r'\[(.*?)\]', text)

    if not bracket_matches:
        return text

    result_text = text
    for placeholder in bracket_matches:
        # Convert placeholder to possible field names
        possible_fields = [
            placeholder.lower(),                      # [company name] -> company name
            placeholder.lower().replace(' ', '_'),    # [company name] -> company_name
            placeholder.lower().replace('-', '_'),    # [company-name] -> company_name
            placeholder.lower().replace(' ', '')      # [company name] -> companyname
        ]

        # Try to find matching field in invoice data
        matched_field = None
        matched_value = None

        # Direct mapping for common placeholders
        direct_mappings = {
            "company name": "company_name",
            "company slogan": "company_slogan",
            "company address": "company_address",
            "company city": "company_city",
            "company state": "company_state",
            "company zip": "company_zip",
            "company phone": "company_phone",
            "company email": "company_email",
            "company website": "company_website",

            "invoice number": "invoice_number",
            "invoice date": "invoice_date",
            "due date": "due_date",
            "payment terms": "payment_terms",

            "client name": "client_name",
            "client company": "client_company",
            "client address": "client_address",
            "client city": "client_city",
            "client state": "client_state",
            "client zip": "client_zip",
            "client phone": "client_phone",
            "client email": "client_email",

            "subtotal": "subtotal",
            "discount amount": "discount_amount",
            "tax amount": "tax_amount",
            "total amount": "total_amount",
            "total": "total_amount"
        }

        # Check for direct mapping first
        placeholder_lower = placeholder.lower()
        if placeholder_lower in direct_mappings:
            mapped_field = direct_mappings[placeholder_lower]
            if mapped_field in invoice_data:
                matched_field = mapped_field
                matched_value = invoice_data[mapped_field]
                logger.info(f"Direct mapping for '{placeholder}' to '{mapped_field}': {matched_value}")

        # If no direct match, try the possible field variations
        if not matched_field:
            for field_name in possible_fields:
                if field_name in invoice_data:
                    matched_field = field_name
                    matched_value = invoice_data[field_name]
                    logger.info(f"Found exact field match for '{placeholder}': {matched_field}")
                    break

            # If still no exact match found, try partial matches
            if not matched_field:
                for key in invoice_data:
                    if (key.lower() in possible_fields or
                        any(field.lower() in key.lower() for field in possible_fields)):
                        matched_field = key
                        matched_value = invoice_data[key]
                        logger.info(f"Found partial field match for '{placeholder}': {matched_field}")
                        break

        # If we found a match, replace the placeholder
        if matched_field and matched_value is not None:
            # Format value based on field type
            formatted_value = format_value(matched_value, matched_field)
            # Replace in text
            result_text = result_text.replace(f"[{placeholder}]", formatted_value)
            logger.info(f"Replaced bracket placeholder '[{placeholder}]' with '{formatted_value}'")
        else:
            logger.warning(f"No match found for bracket placeholder '[{placeholder}]'")

    return result_text

# Testing function
def test_template_selection():
    """Test the template selection logic with some sample data."""
    # Define some test cases
    test_cases = [
        {
            "description": "Service invoice with minimal data",
            "data": {
                "invoice_type": "service",
                "vendor": "Acme Services",
                "items": [
                    {
                        "description": "Consulting hours",
                        "quantity": 10,
                        "unit_price": 150,
                        "total_price": 1500
                    }
                ],
                "total_amount": 1500
            }
        },
        {
            "description": "Time and materials invoice",
            "data": {
                "invoice_type": "time_materials",
                "company_name": "Construction Co",
                "client_name": "Homeowner",
                "items": [
                    {
                        "description": "Labor hours",
                        "quantity": 8,
                        "unit_price": 75,
                        "total_price": 600
                    },
                    {
                        "description": "Materials - Lumber",
                        "quantity": 1,
                        "unit_price": 350,
                        "total_price": 350
                    }
                ],
                "total_amount": 950
            }
        }
    ]

    # Set up logging
    logging.basicConfig(level=logging.INFO)

    # Test each case
    for i, test in enumerate(test_cases, 1):
        print(f"\nTest {i}: {test['description']}")
        template_result = select_template(test["data"])
        print(f"Selected: {template_result['template_type']} ({template_result['template_info']['name']})")
        print(f"Template path: {template_result['template_path']}")

        if template_result['missing_fields']:
            print(f"Missing fields: {', '.join(template_result['missing_fields'])}")
        else:
            print("All required fields present")

if __name__ == "__main__":
    test_template_selection()